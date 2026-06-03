"""Экспорт ПСД-анализа в Excel/DOCX.

Используется как миксин в PsdAnalystService. Обращается к helpers
`self._build_dvc_maps` и `self._calc_min_dvc` (определены в главном классе).
"""
import os
import re
from collections import defaultdict
from datetime import datetime
from decimal import Decimal
from typing import Any, Dict, Optional

import pandas as pd
from sqlalchemy import func
from sqlalchemy.orm import Session

from ..models.models import (
    Agsk, AgskEnstruMatch, ExternalDocument, PsdDocumentItem,
    PsdItemSupplierSelection, Reestr_KTP, User,
)


class PsdExportMixin:
    def export_matches_to_excel(self, db: Session, format_type: str = "full"):
        matches = db.query(AgskEnstruMatch).filter(
            AgskEnstruMatch.is_active == True,
            AgskEnstruMatch.is_approved == True,
        ).order_by(AgskEnstruMatch.agsk_code).all()
        data = [{
            "Код АГСК": m.agsk_code,
            "Код ЕНС ТРУ": m.enstru_code,
            "Дата сопоставления": m.created_at.strftime('%d.%m.%Y') if m.created_at else None,
            "Дата утверждения": m.approved_at.strftime('%d.%m.%Y') if m.approved_at else None,
        } for m in matches]
        path = f"/tmp/export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        pd.DataFrame(data).to_excel(path, index=False)
        return {"file_path": path}

    def export_full_analysis_report(self, db: Session, doc_id: int) -> Optional[str]:
        import openpyxl
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter

        q = db.query(PsdDocumentItem, Agsk).outerjoin(Agsk, PsdDocumentItem.code_sn == Agsk.code)
        q = q.filter(PsdDocumentItem.document_id == doc_id)
        items_data = q.order_by(PsdDocumentItem.id).all()

        # Deduplicate by item.id — outerjoin with Agsk can produce duplicate rows
        # if code_sn matches multiple Agsk records, which inflates averages
        _seen_item_ids: set = set()
        _deduped: list = []
        for _it, _ag in items_data:
            if _it.id not in _seen_item_ids:
                _seen_item_ids.add(_it.id)
                _deduped.append((_it, _ag))
        items_data = _deduped

        if not items_data:
            return None

        doc = db.query(ExternalDocument).filter(ExternalDocument.id == doc_id).first()
        agsk_codes = list({item.code_sn for item, _ in items_data if item.code_sn})
        enstru_codes_set = list({item.enstru_code for item, _ in items_data if item.enstru_code})

        direct_ktp_map, group_ktp_map, suppliers_by_enstru, agsk_all_map = self._build_dvc_maps(
            db, agsk_codes, enstru_codes_set
        )

        wb = openpyxl.Workbook()
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="1B5E20", end_color="1B5E20", fill_type="solid")
        sub_header_fill = PatternFill(start_color="E8F5E9", end_color="E8F5E9", fill_type="solid")
        border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )

        type_letters = {'GOODS': 'Т', 'WORKS': 'Р', 'SERVICES': 'У', 'OTHER': 'П', 'BALANCE': 'Б'}
        type_row_fills = {
            'WORKS': PatternFill(start_color="EBF5FB", end_color="EBF5FB", fill_type="solid"),
            'SERVICES': PatternFill(start_color="EAF7EA", end_color="EAF7EA", fill_type="solid"),
            'OTHER': PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid"),
            'BALANCE': PatternFill(start_color="FEF9E7", end_color="FEF9E7", fill_type="solid"),
        }

        # Группируем позиции по типу
        grouped = defaultdict(list)
        for item, agsk in items_data:
            itype = item.item_type or 'GOODS'
            grouped[itype].append((item, agsk))

        # ── ЛИСТ 1: СМЕТА ────────────────────────────────────────────────────
        ws1 = wb.active
        ws1.title = "Смета"

        ws1.merge_cells('A1:K1')
        ws1['A1'] = "СМЕТА ЗАКУПОК"
        ws1['A1'].font = Font(size=16, bold=True)
        ws1['A1'].alignment = Alignment(horizontal='center')
        if doc:
            ws1.merge_cells('A2:K2')
            ws1['A2'] = f"Наименование проекта: {doc.bank_name}"
            ws1['A2'].font = Font(bold=True, size=12)

        columns1 = [
            "№", "Код ЕНС ТРУ", "Наименование закупаемых товаров услуг работ",
            "Единица измерения(МКЕИ)", "Количество, объём", "Цена за единицу тенге(без НДС)",
            "Сумма планируемая для закупок ТРУ", "КОД АГСК", "КТП", "ВЦ %", "Сумма ВЦ тенге без НДС"
        ]

        def write_header(ws, row):
            for col_idx, col_name in enumerate(columns1, 1):
                cell = ws.cell(row=row, column=col_idx, value=col_name)
                cell.font = header_font
                cell.fill = header_fill
                cell.border = border
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            ws.row_dimensions[row].height = 45
            return row + 1

        def fill_section(ws, curr_row, section_title, section_items, is_first):
            if not section_items:
                return curr_row, Decimal('0'), Decimal('0'), Decimal('0'), 0

            if is_first:
                curr_row = write_header(ws, curr_row)

            # Заголовок секции
            ws.merge_cells(f'A{curr_row}:K{curr_row}')
            c = ws.cell(row=curr_row, column=1, value=section_title)
            c.font = Font(bold=True, size=12)
            c.fill = sub_header_fill
            curr_row += 1

            section_total = Decimal('0')
            section_vc = Decimal('0')
            dvc_pct_sum = Decimal('0')
            dvc_pct_count = 0

            for idx, (item, agsk) in enumerate(section_items, 1):
                itype = item.item_type or 'GOODS'
                num_label = f"{idx} {type_letters.get(itype, 'П')}"
                dvc_p = self._calc_min_dvc(item, direct_ktp_map, group_ktp_map, suppliers_by_enstru, agsk_all_map)
                i_total = Decimal(str(item.total_amount or 0))
                vc_a = i_total * (dvc_p / 100)
                section_total += i_total
                section_vc += vc_a
                dvc_pct_sum += dvc_p
                dvc_pct_count += 1

                row_data = [
                    num_label,
                    item.enstru_code or "",
                    item.name or "",
                    item.unit or "",
                    float(item.volume or 0),
                    float(item.price or 0),
                    float(i_total),
                    item.code_sn or "",
                    "Да" if item.match_type != 'none' else ("100%" if itype in ('WORKS', 'SERVICES') else "Нет"),
                    float(dvc_p),
                    float(vc_a)
                ]

                row_fill = type_row_fills.get(itype)
                for c_idx, val in enumerate(row_data, 1):
                    cell = ws.cell(row=curr_row, column=c_idx, value=val)
                    cell.border = border
                    if c_idx in [5, 6, 7, 10, 11]:
                        cell.number_format = '#,##0.00'
                    if row_fill:
                        cell.fill = row_fill
                curr_row += 1

            # Итог секции — простое арифметическое среднее ВЦ% (все позиции, включая 0)
            avg_dvc_sec = dvc_pct_sum / dvc_pct_count if dvc_pct_count > 0 else Decimal('0')
            ws.merge_cells(f'A{curr_row}:F{curr_row}')
            c = ws.cell(row=curr_row, column=1, value=f"Итого по {section_title.lower()}:")
            c.font = Font(bold=True)
            c.alignment = Alignment(horizontal='right')
            c7 = ws.cell(row=curr_row, column=7, value=float(section_total))
            c7.font = Font(bold=True)
            c7.number_format = '#,##0.00'
            c10 = ws.cell(row=curr_row, column=10, value=float(avg_dvc_sec))
            c10.font = Font(bold=True)
            c10.number_format = '0.00'
            c11 = ws.cell(row=curr_row, column=11, value=float(section_vc))
            c11.font = Font(bold=True)
            c11.number_format = '#,##0.00'
            curr_row += 2  # отступ после секции

            return curr_row, section_total, section_vc, dvc_pct_sum, dvc_pct_count

        curr_row = 6
        total_sum = Decimal('0')
        total_vc_sum = Decimal('0')
        total_dvc_pct_sum = Decimal('0')
        total_dvc_pct_count = 0

        section_configs = [
            ('1. Товары (Т)', 'GOODS'),
            ('2. Работы (Р)', 'WORKS'),
            ('3. Услуги (У)', 'SERVICES'),
            ('4. Прочее (П)', 'OTHER'),
        ]

        # Build item_id → num_label map so sheet 2 uses the same numbering as sheet 1
        item_num_labels: dict = {}
        for _, itype_key in section_configs:
            tl = type_letters.get(itype_key, 'П')
            for sec_idx, (sec_item, _) in enumerate(grouped.get(itype_key, []), 1):
                item_num_labels[sec_item.id] = f"{sec_idx} {tl}"

        is_first = True
        for section_title, itype in section_configs:
            items_in_section = grouped.get(itype, [])
            if not items_in_section:
                continue
            curr_row, sec_total, sec_vc, sec_dvc_sum, sec_dvc_count = fill_section(
                ws1, curr_row, section_title, items_in_section, is_first
            )
            total_sum += sec_total
            total_vc_sum += sec_vc
            total_dvc_pct_sum += sec_dvc_sum
            total_dvc_pct_count += sec_dvc_count
            is_first = False

        # Общий итог — простое арифметическое среднее ВЦ% по всем позициям
        avg_total_dvc = total_dvc_pct_sum / total_dvc_pct_count if total_dvc_pct_count > 0 else Decimal('0')
        ws1.merge_cells(f'A{curr_row}:F{curr_row}')
        ws1.cell(row=curr_row, column=1, value="ИТОГО ПО СМЕТЕ:").font = Font(bold=True, size=11)
        ws1.cell(row=curr_row, column=1).alignment = Alignment(horizontal='right')
        ws1.cell(row=curr_row, column=7, value=float(total_sum)).font = Font(bold=True, size=11)
        ws1.cell(row=curr_row, column=7).number_format = '#,##0.00'
        c10_total = ws1.cell(row=curr_row, column=10, value=float(avg_total_dvc))
        c10_total.font = Font(bold=True, size=11)
        c10_total.number_format = '0.00'
        ws1.cell(row=curr_row, column=11, value=float(total_vc_sum)).font = Font(bold=True, size=11)
        ws1.cell(row=curr_row, column=11).number_format = '#,##0.00'

        # ── ЛИСТ 2: АНАЛИЗ ПСД — только GOODS ───────────────────────────────
        ws2 = wb.create_sheet("Анализ ПСД")
        columns2 = [
            "№ позиции", "Наименование позиции", "Код АГСК", "Полное название АГСК",
            "Ед. изм.", "Объем", "Цена", "Сумма",
            "Источник", "Производитель из РКТП", "БИН",
            "Наименование товара из РКТП", "ВЦ%", "Адрес", "Регион",
            "Коды АГСК из Реестра КТП"
        ]
        for col_idx, col_name in enumerate(columns2, 1):
            cell = ws2.cell(row=1, column=col_idx, value=col_name)
            cell.font = header_font
            cell.fill = PatternFill("solid", fgColor="2C3E50")
            cell.border = border
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

        FILL_MAP = {
            "Выбор аналитика": PatternFill("solid", fgColor="D5F5E3"),   # зелёный
            "Авто-подбор (КТП)": PatternFill("solid", fgColor="D6EAF8"), # голубой
        }

        # Предзагружаем активные выборы поставщиков для всех позиций GOODS
        goods_item_ids = [item.id for item, _ in grouped.get('GOODS', [])]
        active_selections_by_item: Dict[int, list] = defaultdict(list)
        if goods_item_ids:
            all_active_sels = db.query(PsdItemSupplierSelection).filter(
                PsdItemSupplierSelection.item_id.in_(goods_item_ids),
                PsdItemSupplierSelection.status == 'active',
                PsdItemSupplierSelection.is_active == True,
            ).all()
            for sel in all_active_sels:
                active_selections_by_item[sel.item_id].append(sel)

        # Предзагружаем KTP-записи для выборов аналитика (нужны адрес, регион, АГСК-коды)
        all_ktp_ids = list({
            sel.ktp_id
            for sels in active_selections_by_item.values()
            for sel in sels
            if sel.ktp_id
        })
        ktp_by_id: Dict[int, Any] = {}
        if all_ktp_ids:
            ktp_recs = db.query(Reestr_KTP).filter(Reestr_KTP.id.in_(all_ktp_ids)).all()
            ktp_by_id = {r.id: r for r in ktp_recs}

        curr_row2 = 2

        # Только GOODS (с сопоставлением ИЛИ с выбором аналитика ИЛИ с АГСК в реестре КТП)
        for item, agsk in grouped.get('GOODS', []):
            has_selection = bool(active_selections_by_item.get(item.id))
            has_agsk_match = bool(item.code_sn and agsk_all_map.get(item.code_sn))
            if not has_selection and not has_agsk_match:
                continue  # нечего показывать — пропускаем

            num_label = item_num_labels.get(item.id, "? Т")
            base = [
                num_label,
                item.name,
                item.code_sn,
                agsk.full_name if agsk else "",
                item.unit,
                float(item.volume or 0),
                float(item.price or 0),
                float(item.total_amount or 0),
            ]

            rows_to_add = []

            # Приоритет 1: выборы аналитика (из psd_item_supplier_selections со статусом active)
            if has_selection:
                for sel in active_selections_by_item[item.id]:
                    dvc = float(sel.dvc_percent) if sel.dvc_percent else 0
                    # Подтягиваем адрес / регион / АГСК-коды из реестра КТП по ktp_id
                    ktp_rec = ktp_by_id.get(sel.ktp_id) if sel.ktp_id else None
                    rows_to_add.append(base + [
                        "Выбор аналитика",
                        sel.supplier_name or "—",
                        sel.supplier_bin or "—",
                        sel.supplier_product or "—",
                        dvc,
                        ktp_rec.production_address if ktp_rec else "",
                        ktp_rec.region_kato if ktp_rec else "",
                        ", ".join(ktp_rec.agsk3_codes) if (ktp_rec and ktp_rec.agsk3_codes) else "",
                    ])

            elif has_agsk_match:
                # Нет выбора аналитика → показываем ТОЛЬКО поставщиков с совпадающим АГСК
                for ktp_r in agsk_all_map[item.code_sn]:
                    dvc = float(
                        re.sub(r'[^0-9.]', '', ktp_r.dvc_percent).replace(',', '.')
                    ) if ktp_r.dvc_percent else 0
                    if dvc <= 0:
                        continue
                    rows_to_add.append(base + [
                        "Авто-подбор (КТП)",
                        ktp_r.company_name, ktp_r.bin_iin, ktp_r.product_name,
                        dvc, ktp_r.production_address, ktp_r.region_kato,
                        ", ".join(ktp_r.agsk3_codes) if ktp_r.agsk3_codes else "",
                    ])

            if not rows_to_add:
                continue

            for r_data in rows_to_add:
                src = r_data[8]
                fill = FILL_MAP.get(src, PatternFill())
                for c_idx, val in enumerate(r_data, 1):
                    cell = ws2.cell(row=curr_row2, column=c_idx, value=val)
                    cell.fill = fill
                    cell.border = border
                    if c_idx in [6, 7, 8, 13]:
                        cell.number_format = '#,##0.00'
                ws2.cell(row=curr_row2, column=12).font = Font(bold=True)
                curr_row2 += 1

        # Автоширина колонок
        for ws in [ws1, ws2]:
            for i, col in enumerate(ws.columns, 1):
                m_l = 0
                for cell in col:
                    try:
                        if cell.value and len(str(cell.value)) > m_l:
                            m_l = len(str(cell.value))
                    except:
                        pass
                ws.column_dimensions[get_column_letter(i)].width = min(m_l + 2, 40)

        os.makedirs("/tmp", exist_ok=True)
        file_path = f"/tmp/psd_analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        wb.save(file_path)
        return file_path

    def generate_psd_conclusion_docx(self, db: Session, doc_id: int, current_user: User) -> Optional[str]:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc_entity = db.query(ExternalDocument).filter(ExternalDocument.id == doc_id).first()
        if not doc_entity:
            return None

        total_items = db.query(PsdDocumentItem).filter(PsdDocumentItem.document_id == doc_id).count()
        # «Обработанными» считаются позиции с активным выбором поставщика
        # (match_type='suggested' — только подсказка, не обработка)
        matched_items = db.query(PsdDocumentItem).filter(
            PsdDocumentItem.document_id == doc_id,
            PsdDocumentItem.match_type == 'manual'
        ).count()
        total_amount = db.query(func.sum(PsdDocumentItem.total_amount)).filter(
            PsdDocumentItem.document_id == doc_id
        ).scalar() or 0

        items = db.query(PsdDocumentItem).filter(PsdDocumentItem.document_id == doc_id).all()

        agsk_codes_doc = list({it.code_sn for it in items if it.code_sn})
        enstru_codes_doc = list({it.enstru_code for it in items if it.enstru_code})
        direct_ktp_map, group_ktp_map, suppliers_map, agsk_all_map_doc = self._build_dvc_maps(
            db, agsk_codes_doc, enstru_codes_doc
        )

        goods_amount = Decimal('0')
        works_amount = Decimal('0')
        services_amount = Decimal('0')
        other_amount = Decimal('0')
        goods_vc = Decimal('0')
        works_vc = Decimal('0')
        services_vc = Decimal('0')
        total_vc_amount = Decimal('0')
        goods_dvc_sum = Decimal('0'); goods_dvc_count = 0
        works_dvc_sum = Decimal('0'); works_dvc_count = 0
        services_dvc_sum = Decimal('0'); services_dvc_count = 0
        total_dvc_sum = Decimal('0'); total_dvc_count = 0

        for it in items:
            itype = it.item_type or 'GOODS'
            i_total = Decimal(str(it.total_amount or 0))
            dvc = self._calc_min_dvc(it, direct_ktp_map, group_ktp_map, suppliers_map, agsk_all_map_doc)
            vc_part = i_total * (dvc / 100)
            total_dvc_sum += dvc
            total_dvc_count += 1
            if itype == 'GOODS':
                goods_amount += i_total
                goods_vc += vc_part
                goods_dvc_sum += dvc
                goods_dvc_count += 1
            elif itype == 'WORKS':
                works_amount += i_total
                works_vc += vc_part
                works_dvc_sum += dvc
                works_dvc_count += 1
            elif itype == 'SERVICES':
                services_amount += i_total
                services_vc += vc_part
                services_dvc_sum += dvc
                services_dvc_count += 1
            else:
                other_amount += i_total
            total_vc_amount += vc_part

        # Простое арифметическое среднее (все позиции включая 0)
        avg_vc_percent = total_dvc_sum / total_dvc_count if total_dvc_count > 0 else Decimal('0')
        goods_avg_dvc = goods_dvc_sum / goods_dvc_count if goods_dvc_count > 0 else Decimal('0')
        works_avg_dvc = works_dvc_sum / works_dvc_count if works_dvc_count > 0 else Decimal('0')
        services_avg_dvc = services_dvc_sum / services_dvc_count if services_dvc_count > 0 else Decimal('0')

        doc = Document()
        title = doc.add_heading('ЗАКЛЮЧЕНИЕ АНАЛИТИКА ДРВЦ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        p.add_run('Дата формирования: ').bold = True
        p.add_run(datetime.now().strftime('%d.%m.%Y %H:%M'))

        p = doc.add_paragraph()
        p.add_run('Аналитик: ').bold = True
        p.add_run(f"{current_user.full_name}")

        p = doc.add_paragraph()
        p.add_run('Наименование Банка: ').bold = True
        p.add_run(f"{doc_entity.bank_name}")

        doc.add_heading('Результаты анализа', level=1)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Показатель'
        hdr_cells[1].text = 'Значение'

        data = [
            ('Общее количество позиций', str(total_items)),
            ('Сопоставлено позиций (Товары)', str(matched_items)),
            ('Процент сопоставления', f"{(matched_items / total_items * 100):.1f}%" if total_items > 0 else "0%"),
            ('', ''),
            ('Сумма Товары (Т)', f"{float(goods_amount):,.2f} тенге"),
            ('Средний ВЦ% по Товарам (Т)', f"{float(goods_avg_dvc):.2f}%"),
            ('Сумма Работы (Р)', f"{float(works_amount):,.2f} тенге"),
            ('Средний ВЦ% по Работам (Р)', f"{float(works_avg_dvc):.2f}%"),
            ('Сумма Услуги (У)', f"{float(services_amount):,.2f} тенге"),
            ('Средний ВЦ% по Услугам (У)', f"{float(services_avg_dvc):.2f}%"),
        ]
        if other_amount > 0:
            data.append(('Сумма Прочее (П)', f"{float(other_amount):,.2f} тенге"))
        data += [
            ('Общая сумма проекта', f"{float(total_amount):,.2f} тенге"),
            ('', ''),
            ('Прогнозируемая сумма ВЦ', f"{float(total_vc_amount):,.2f} тенге"),
            ('Средний ВЦ% (общий)', f"{float(avg_vc_percent):.2f}%"),
        ]

        for label, value in data:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value

        doc.add_paragraph()
        doc.add_heading('Выводы', level=1)
        conclusion_text = (
            f"В ходе анализа проектно-сметной документации '{doc_entity.bank_name}' "
            f"было обработано {total_items} позиций. "
            f"Уровень внутристрановой ценности (ВЦ) по проекту оценивается в {float(avg_vc_percent):.2f}%. "
        )
        doc.add_paragraph(conclusion_text)

        if doc_entity.analyst_comment:
            doc.add_paragraph()
            doc.add_heading('Дополнительный комментарий аналитика', level=1)
            doc.add_paragraph(doc_entity.analyst_comment)

        doc.add_paragraph().add_run('\n\n__________________________ / Подпись /').italic = True

        os.makedirs("/tmp", exist_ok=True)
        file_path = f"/tmp/conclusion_{doc_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(file_path)
        return file_path

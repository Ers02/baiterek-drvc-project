from decimal import Decimal

from sqlalchemy.orm import Session, joinedload
from sqlalchemy import func, or_, and_, String, desc, text, cast, case
from typing import List, Optional, Dict, Any, Literal, Union
from datetime import datetime, timezone, timedelta
import pandas as pd
import os
import re
import zipfile
import shutil
import uuid
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..models.models import (
    PsdDocumentItem, AgskEnstruMatch,
    ExternalDocument, Reestr_KTP, Agsk, Enstru, PsdAnalysisSession, User, UserRole,
    Oked, Kpved, Tnved
)
from ..utils.text_utils import tokenize, score_pair

SearchMode = Literal["all", "agsk", "name"]


class PsdAnalystService:

    @staticmethod
    def _build_dvc_maps(db: Session, agsk_codes: list, enstru_codes: list):
        """Bulk-build DVC lookup maps. Returns (direct_ktp_map, group_ktp_map, suppliers_map)."""
        suppliers_map: Dict[str, list] = {}
        for code in enstru_codes:
            suppliers_map[code] = db.query(Reestr_KTP).filter(
                Reestr_KTP.enstru_codes.contains([code])
            ).all()

        direct_ktp_map: Dict[str, Any] = {}
        group_ktp_map: Dict[str, Any] = {}
        if agsk_codes:
            agsk_set = set(agsk_codes)
            ktp_bulk = db.query(Reestr_KTP).filter(
                or_(*[Reestr_KTP.agsk3_codes.contains([c]) for c in agsk_codes])
            ).order_by(Reestr_KTP.id).all()
            for ktp in ktp_bulk:
                for ac in (ktp.agsk3_codes or []):
                    if ac in agsk_set and ac not in direct_ktp_map:
                        direct_ktp_map[ac] = ktp
            missing = [c for c in agsk_codes if c not in direct_ktp_map and len(c) >= 10]
            seen_parents: set = set()
            for code in missing:
                parent = code[:10]
                if parent in seen_parents:
                    continue
                seen_parents.add(parent)
                rg = db.query(Reestr_KTP).filter(
                    text("EXISTS (SELECT 1 FROM jsonb_array_elements_text(agsk3_codes) AS elem WHERE elem LIKE :prefix)")
                ).params(prefix=f"{parent}%").order_by(Reestr_KTP.id).first()
                if rg:
                    for c2 in missing:
                        if c2[:10] == parent:
                            group_ktp_map[c2] = rg

        return direct_ktp_map, group_ktp_map, suppliers_map

    @staticmethod
    def _calc_min_dvc(item, direct_ktp_map, group_ktp_map, suppliers_map) -> Decimal:
        """Return minimum DVC% for a PsdDocumentItem using pre-built lookup maps."""
        itype = item.item_type or 'GOODS'
        if itype in ('WORKS', 'SERVICES'):
            return Decimal('100')
        all_vals = []
        for ktp_r in [direct_ktp_map.get(item.code_sn), group_ktp_map.get(item.code_sn)]:
            if ktp_r and ktp_r.dvc_percent:
                try:
                    v = float(re.sub(r'[^0-9.]', '', ktp_r.dvc_percent).replace(',', '.'))
                    if v > 0:
                        all_vals.append(v)
                except Exception:
                    pass
        if item.enstru_code:
            for s in suppliers_map.get(item.enstru_code, []):
                if s.dvc_percent:
                    try:
                        v = float(re.sub(r'[^0-9.]', '', s.dvc_percent).replace(',', '.'))
                        if v > 0:
                            all_vals.append(v)
                    except Exception:
                        pass
        return Decimal(str(min(all_vals))) if all_vals else Decimal('0')

    @staticmethod
    def _calculate_deadline(start_date: datetime, business_days: int) -> datetime:
        current_date = start_date
        added_days = 0
        while added_days < business_days:
            current_date += timedelta(days=1)
            if current_date.weekday() < 5:
                added_days += 1
        return current_date

    def assign_to_analyst(self, db: Session, doc_id: int, analyst_id: int, deadline_days: int):
        doc = db.query(ExternalDocument).filter(ExternalDocument.id == doc_id).first()
        if not doc:
            raise ValueError("Документ не найден")
        doc.assigned_to = analyst_id
        doc.assigned_at = func.now()
        doc.deadline_days = deadline_days
        doc.deadline_at = self._calculate_deadline(datetime.now(), deadline_days)
        doc.status = "ASSIGNED_TO_ANALYST"
        if doc.external_id and doc.bank_name:
            related_docs = db.query(ExternalDocument).filter(
                ExternalDocument.bank_name == doc.bank_name,
                ExternalDocument.external_id == doc.external_id,
                ExternalDocument.id != doc_id,
                ExternalDocument.assigned_to.is_(None)
            ).all()
            for related in related_docs:
                related.assigned_to = analyst_id
                related.assigned_at = func.now()
                related.deadline_days = deadline_days
                related.deadline_at = doc.deadline_at
                related.status = "ASSIGNED_TO_ANALYST"
        db.commit()
        return doc

    def submit_for_approval(self, db: Session, doc_id: int):
        doc = db.query(ExternalDocument).filter(ExternalDocument.id == doc_id).first()
        if not doc:
            raise ValueError("Документ не найден")
        doc.status = "FOR_APPROVAL"
        db.commit()
        return doc

    def approve_document(self, db: Session, doc_id: int, director: User):
        doc = db.query(ExternalDocument).filter(ExternalDocument.id == doc_id).first()
        if not doc:
            raise ValueError("Документ не найден")
        doc.status = "APPROVED"
        excel_path = self.export_full_analysis_report(db, doc_id)
        docx_path = self.generate_psd_conclusion_docx(db, doc_id, doc.assigned_user or director)
        zip_filename = f"result_{doc_id}_{uuid.uuid4().hex[:8]}.zip"
        zip_path = os.path.join("/tmp", zip_filename)
        with zipfile.ZipFile(zip_path, 'w') as zf:
            if excel_path and os.path.exists(excel_path):
                zf.write(excel_path, arcname=f"report_{doc_id}.xlsx")
            if docx_path and os.path.exists(docx_path):
                zf.write(docx_path, arcname=f"conclusion_{doc_id}.docx")
        doc.result_file_path = zip_path
        doc.status = "COMPLETED"
        doc.completed_at = func.now()
        db.commit()
        return doc

    def reject_document(self, db: Session, doc_id: int, comment: str):
        doc = db.query(ExternalDocument).filter(ExternalDocument.id == doc_id).first()
        if not doc:
            raise ValueError("Документ не найден")
        doc.status = "REJECTED_BY_DIRECTOR"
        doc.director_comment = comment
        db.commit()
        return doc

    def delegate_authority(self, db: Session, from_user_id: int, to_user_id: int, days: int):
        user = db.query(User).filter(User.id == from_user_id).first()
        if not user:
            raise ValueError("Пользователь не найден")
        user.delegated_to_id = to_user_id
        user.delegation_start = func.now()
        user.delegation_end = datetime.now() + timedelta(days=days)
        db.commit()
        return user

    @staticmethod
    def _reestr_to_result(r: Reestr_KTP, enstru_code: str, enstru_name: str) -> Dict:
        dvc_val = float(re.sub(r'[^0-9.]', '', r.dvc_percent)) if r.dvc_percent else 0
        return {
            "ktp_id": r.id,
            "enstru_code": enstru_code,
            "enstru_name": enstru_name or "—",
            "company": r.company_name,
            "bin": r.bin_iin,
            "product": r.product_name,
            "dvc_percent": dvc_val,
            "address": r.production_address,
            "registry_date": r.registry_inclusion_date.strftime('%d.%m.%Y') if r.registry_inclusion_date else None,
            "source": "reestr_ktp",
            "region": r.region_kato,
            "agsk3_codes": r.agsk3_codes or [],
            "agsk3_names": r.agsk3_names or [],
            "oked_codes": r.oked_codes or [],
            "oked_names": r.oked_names or [],
            "kpved_codes": r.kpved_codes or [],
            "kpved_names": r.kpved_names or [],
            "tnved_codes": r.tnved_codes or [],
        }

    def _run_auto_matching_for_document(self, db: Session, doc_id: int):
        from .agsk_enstru_matcher import AgskEnstruMatcher
        items = db.query(PsdDocumentItem).filter(PsdDocumentItem.document_id == doc_id).all()

        # Bulk-load утверждённых сопоставлений из проверенной библиотеки
        agsk_codes = [it.code_sn for it in items if it.code_sn]
        approved_map: Dict[str, AgskEnstruMatch] = {}
        if agsk_codes:
            approved_rows = db.query(AgskEnstruMatch).filter(
                AgskEnstruMatch.agsk_code.in_(agsk_codes),
                AgskEnstruMatch.is_approved == True,
                AgskEnstruMatch.is_active == True,
            ).order_by(AgskEnstruMatch.approved_at.asc()).all()
            # Последнее утверждённое по каждому АГСК (asc → последнее перезаписывает)
            for m in approved_rows:
                approved_map[m.agsk_code] = m

        # Bulk-load имён ЕНСТРУ для утверждённых кодов
        approved_enstru_codes = list({m.enstru_code for m in approved_map.values()})
        enstru_name_map: Dict[str, str] = {}
        if approved_enstru_codes:
            enstru_name_map = {
                e.code: e.name_rus
                for e in db.query(Enstru).filter(Enstru.code.in_(approved_enstru_codes)).all()
            }

        matcher = AgskEnstruMatcher(db)
        for item in items:
            if not item.code_sn:
                continue

            # Приоритет 1: утверждённая библиотека (проверено менеджером)
            approved = approved_map.get(item.code_sn)
            if approved:
                item.enstru_code = approved.enstru_code
                item.enstru_name = enstru_name_map.get(approved.enstru_code)
                item.match_type = "manual"
                item.match_score = 100
                item.match_reason = "Автоматически из утверждённой библиотеки сопоставлений"
                continue

            # Приоритет 2: обычный авто-матчинг
            best = matcher.get_match_for_agsk(item.code_sn)
            if best:
                item.enstru_code = best["enstru_code"]
                item.enstru_name = best["enstru_name"]
                item.match_type = best["match_type"]
                item.match_score = best["score"]
                item.match_reason = best["reason"]
            else:
                item.match_type = "none"
        db.commit()

    def search_enstru_in_reestr(self, db: Session, query: str, limit: int = 20, search_mode: SearchMode = "all") -> List[Dict]:
        q = query.strip()
        if not q or len(q) < 2:
            return []
        base_filter = text("NULLIF(REGEXP_REPLACE(dvc_percent, '[^0-9.]', '', 'g'), '')::numeric > 0")
        if search_mode == "agsk":
            mode_filter = text("agsk3_codes::text ILIKE :agsk_prefix").bindparams(agsk_prefix=f'%"{q}%')
        elif search_mode == "name":
            mode_filter = or_(
                Reestr_KTP.product_name.ilike(f"%{q}%"),
                Reestr_KTP.company_name.ilike(f"%{q}%"),
                cast(Reestr_KTP.enstru_names, String).ilike(f"%{q}%"),
            )
        else:
            mode_filter = or_(
                Reestr_KTP.product_name.ilike(f"%{q}%"),
                Reestr_KTP.company_name.ilike(f"%{q}%"),
                cast(Reestr_KTP.enstru_codes, String).ilike(f"%{q}%"),
                cast(Reestr_KTP.enstru_names, String).ilike(f"%{q}%"),
                cast(Reestr_KTP.agsk3_codes, String).ilike(f"%{q}%"),
                cast(Reestr_KTP.oked_codes, String).ilike(f"%{q}%"),
                cast(Reestr_KTP.kpved_codes, String).ilike(f"%{q}%"),
                cast(Reestr_KTP.tnved_codes, String).ilike(f"%{q}%"),
                cast(Reestr_KTP.oked_names, String).ilike(f"%{q}%"),
                cast(Reestr_KTP.kpved_names, String).ilike(f"%{q}%"),
            )
        sort_expr = case((Reestr_KTP.product_name.ilike(f"%{q}%"), 1), else_=2)
        rows = db.query(Reestr_KTP).filter(and_(base_filter, mode_filter)).order_by(sort_expr).all()

        # Bulk-lookup всех кодов из наших справочников
        all_enstru_codes: set = set()
        all_agsk3_codes: set = set()
        all_oked_codes: set = set()
        all_kpved_codes: set = set()
        all_tnved_codes: set = set()
        for r in rows:
            for c in (r.enstru_codes or []):
                if c: all_enstru_codes.add(str(c))
            for c in (r.agsk3_codes or []):
                if c: all_agsk3_codes.add(str(c))
            for c in (r.oked_codes or []):
                if c: all_oked_codes.add(str(c))
            for c in (r.kpved_codes or []):
                if c: all_kpved_codes.add(str(c))
            for c in (r.tnved_codes or []):
                if c: all_tnved_codes.add(str(c))

        enstru_map: Dict[str, Any] = {e.code: e for e in db.query(Enstru).filter(Enstru.code.in_(all_enstru_codes)).all()} if all_enstru_codes else {}
        agsk_map: Dict[str, Any] = {a.code: a for a in db.query(Agsk).filter(Agsk.code.in_(all_agsk3_codes)).all()} if all_agsk3_codes else {}
        oked_map: Dict[str, Any] = {o.code: o for o in db.query(Oked).filter(Oked.code.in_(all_oked_codes)).all()} if all_oked_codes else {}
        kpved_map: Dict[str, Any] = {k.code: k for k in db.query(Kpved).filter(Kpved.code.in_(all_kpved_codes)).all()} if all_kpved_codes else {}
        tnved_map: Dict[str, Any] = {t.code: t for t in db.query(Tnved).filter(Tnved.code.in_(all_tnved_codes)).all()} if all_tnved_codes else {}

        matches: List[Dict] = []
        seen: set = set()
        for r in rows:
            codes = r.enstru_codes or []
            names = r.enstru_names or []
            if not codes:
                codes = [""]
                names = [""]

            dvc_val = float(re.sub(r'[^0-9.]', '', r.dvc_percent)) if r.dvc_percent else 0
            agsk3_codes = r.agsk3_codes or []
            oked_codes = r.oked_codes or []
            kpved_codes = r.kpved_codes or []
            tnved_codes = r.tnved_codes or []

            for i in range(max(len(codes), len(names))):
                c = codes[i] if i < len(codes) else None
                n = names[i] if i < len(names) else None
                if search_mode == "all":
                    q_lower = q.lower()
                    relevant = (
                        (c and q_lower in str(c).lower()) or (n and q_lower in n.lower())
                        or (r.product_name and q_lower in r.product_name.lower())
                        or (r.company_name and q_lower in r.company_name.lower())
                        or (r.agsk3_codes and q_lower in str(r.agsk3_codes).lower())
                        or (r.oked_codes and q_lower in str(r.oked_codes).lower())
                        or (r.kpved_codes and q_lower in str(r.kpved_codes).lower())
                        or (r.tnved_codes and q_lower in str(r.tnved_codes).lower())
                        or (r.oked_names and q_lower in str(r.oked_names).lower())
                        or (r.kpved_names and q_lower in str(r.kpved_names).lower())
                    )
                    if not relevant:
                        continue
                enstru_obj = enstru_map.get(c) if c else None
                key = (c or "", r.company_name, r.product_name, r.id)
                if key not in seen:
                    matches.append({
                        "ktp_id": r.id,
                        "enstru_code": c or "",
                        "enstru_name": n or "—",
                        "enstru_name_rus": enstru_obj.name_rus if enstru_obj else None,
                        "enstru_detail_rus": enstru_obj.detail_rus if enstru_obj else None,
                        "enstru_standard": enstru_obj.standard if enstru_obj else None,
                        "company": r.company_name,
                        "bin": r.bin_iin,
                        "product": r.product_name,
                        "dvc_percent": dvc_val,
                        "address": r.production_address,
                        "registry_date": r.registry_inclusion_date.strftime('%d.%m.%Y') if r.registry_inclusion_date else None,
                        "source": "reestr_ktp",
                        "region": r.region_kato,
                        "agsk3_codes": agsk3_codes,
                        "agsk3_names": [agsk_map[c2].name_ru if c2 in agsk_map else "" for c2 in agsk3_codes],
                        "oked_codes": oked_codes,
                        "oked_names": [oked_map[c2].name_ru if c2 in oked_map else "" for c2 in oked_codes],
                        "kpved_codes": kpved_codes,
                        "kpved_names": [kpved_map[c2].name_ru if c2 in kpved_map else "" for c2 in kpved_codes],
                        "tnved_codes": tnved_codes,
                        "tnved_names": [tnved_map[c2].name if c2 in tnved_map else "" for c2 in tnved_codes],
                    })
                    seen.add(key)
            if len(matches) >= limit:
                break
        return matches[:limit]

    def _tmp_delete_me(self, db: Session, agsk_code: str, limit: int = 10) -> List[Dict]:
        clean_agsk = agsk_code.strip()
        recs: List[Dict] = []
        seen: set = set()

        def _add_from_ktp(r: Reestr_KTP, score: int, reason: str):
            codes = r.enstru_codes or []
            names = r.enstru_names or []
            dvc_val = float(re.sub(r'[^0-9.]', '', r.dvc_percent)) if r.dvc_percent else 0
            for c, n in zip(codes, names):
                if (c, r.id) not in seen:
                    recs.append({
                        "enstru_code": c, "enstru_name": n, "score": score, "reason": reason,
                        "ktp_id": r.id, "product": r.product_name, "dvc_percent": dvc_val,
                        "agsk3_codes": r.agsk3_codes or [], "agsk3_names": r.agsk3_names or [],
                    })
                    seen.add((c, r.id))

        for r in db.query(Reestr_KTP).filter(text("agsk3_codes::text ILIKE :q").params(q=f'%"{clean_agsk}"%')).all():
            _add_from_ktp(r, 100, f"Точный код АГСК в КТП ({r.company_name})")
        if len(recs) < limit and len(clean_agsk) >= 7:
            parent = clean_agsk[:7]
            for r in db.query(Reestr_KTP).filter(text("agsk3_codes::text ILIKE :q").params(q=f'"{parent}%"')).all():
                _add_from_ktp(r, 80, f"По родительскому коду {parent} в КТП")
        if len(recs) < limit:
            item = db.query(PsdDocumentItem).filter(PsdDocumentItem.code_sn == clean_agsk).first()
            if item:
                tokens = tokenize(item.name)
                if tokens:
                    for r in db.query(Reestr_KTP).filter(or_(*[Reestr_KTP.product_name.ilike(f"%{t}%") for t in tokens[:2]])).all():
                        codes = r.enstru_codes or []
                        names = r.enstru_names or []
                        dvc_val = float(re.sub(r'[^0-9.]', '', r.dvc_percent)) if r.dvc_percent else 0
                        for c, n in zip(codes, names):
                            if (c, r.id) not in seen:
                                sc, _ = score_pair(
                                    {"full_name": item.name, "name_ru": item.name, "standart": ""},
                                    {"name_rus": n, "detail_rus": r.product_name, "standard": ""},
                                )
                                recs.append({
                                    "enstru_code": c, "enstru_name": n, "score": sc,
                                    "reason": "Похожее название товара в КТП", "ktp_id": r.id,
                                    "product": r.product_name, "dvc_percent": dvc_val,
                                    "agsk3_codes": r.agsk3_codes or [], "agsk3_names": r.agsk3_names or [],
                                })
                                seen.add((c, r.id))

    def get_document_items_with_matches(
            self, db: Session, doc_id: int,
            only_unmatched: bool = False,
            search: Optional[str] = None,
            skip: int = 0, limit: int = 50,
    ):
        query = db.query(PsdDocumentItem).filter(PsdDocumentItem.document_id == doc_id)
        if search:
            q = search.strip()
            query = query.filter(or_(
                PsdDocumentItem.name.ilike(f"%{q}%"),
                PsdDocumentItem.code_sn.ilike(f"%{q}%"),
            ))
        if only_unmatched:
            query = query.filter(
                PsdDocumentItem.match_type == "none",
                PsdDocumentItem.item_type.in_(['GOODS', None]),  # только товары
                (PsdDocumentItem.not_in_ktp_registry == False) | (PsdDocumentItem.not_in_ktp_registry == None)
            )
        query = query.order_by(
            case((PsdDocumentItem.item_type == 'OTHER', 2), else_=0),
            desc(
                (PsdDocumentItem.match_type == "none") &
                ((PsdDocumentItem.not_in_ktp_registry == False) | (PsdDocumentItem.not_in_ktp_registry == None))
            ),
            PsdDocumentItem.match_score.asc(),
            PsdDocumentItem.id.asc()
        )
        total_count = query.count()
        items = query.offset(skip).limit(limit).all()
        agsk_codes = [it.code_sn for it in items if it.code_sn]
        agsk_map = {}
        if agsk_codes:
            agsk_data = db.query(Agsk).filter(Agsk.code.in_(agsk_codes)).all()
            agsk_map = {a.code: a for a in agsk_data}

        # Загружаем ВСЕ активные ручные сопоставления для каждой позиции
        item_ids = [it.id for it in items]
        from collections import defaultdict as _defaultdict
        matches_by_item: Dict[int, list] = _defaultdict(list)
        if item_ids:
            matches_list = db.query(AgskEnstruMatch).filter(
                AgskEnstruMatch.item_id.in_(item_ids),
                AgskEnstruMatch.is_active == True,
            ).order_by(AgskEnstruMatch.id.desc()).all()
            for m in matches_list:
                matches_by_item[m.item_id].append(m)

        # Считаем общее кол-во неутверждённых сопоставлений по документу
        pending_match_count = 0
        if doc_id:
            pending_match_count = db.query(AgskEnstruMatch).filter(
                AgskEnstruMatch.doc_id == doc_id,
                AgskEnstruMatch.is_active == True,
                AgskEnstruMatch.is_approved == False,
            ).count()

        result = []
        for item in items:
            agsk_info = agsk_map.get(item.code_sn)
            item_matches = matches_by_item.get(item.id, [])
            current_manual_matches = []
            for lm in item_matches:
                m_status = "approved" if lm.is_approved else "pending"
                current_manual_matches.append({
                    "id": lm.id,
                    "enstru_code": lm.enstru_code,
                    "status": m_status,
                    "matched_at": lm.matched_at.isoformat() if lm.matched_at else None,
                    "approved_at": lm.approved_at.isoformat() if lm.approved_at else None,
                })

            result.append({
                "id": item.id, "item_id": item.id, "document_id": item.document_id,
                "position_number": item.position_number, "name": item.name, "code_sn": item.code_sn,
                "unit": item.unit,
                "volume": float(item.volume) if item.volume else 0,
                "price": float(item.price) if item.price else 0,
                "total_amount": float(item.total_amount) if item.total_amount else 0,
                "enstru_code": item.enstru_code, "enstru_name": item.enstru_name,
                "match_type": item.match_type, "match_score": item.match_score,
                "match_reason": item.match_reason,
                "not_in_ktp_registry": bool(item.not_in_ktp_registry) if item.not_in_ktp_registry is not None else False,
                "can_edit": True,
                "agsk_name_ru": agsk_info.name_ru if agsk_info else None,
                "agsk_full_name": agsk_info.full_name if agsk_info else None,
                "item_type": item.item_type or "GOODS",
                "current_manual_matches": current_manual_matches,
            })
        return {"items": result, "total": total_count, "skip": skip, "limit": limit, "pending_match_count": pending_match_count}

    # ── Ручные сопоставления с approval-воркфлоу ────────────────────────────

    def save_analyst_match(self, db: Session, item_id: int, enstru_code: str, analyst_id: int) -> AgskEnstruMatch:
        """Аналитик добавляет ручное сопоставление для позиции. Поддерживает несколько ЕНСТРУ на одну позицию."""
        item = db.query(PsdDocumentItem).filter(PsdDocumentItem.id == item_id).first()
        if not item:
            raise ValueError("Позиция документа не найдена")

        # Проверяем что такой enstru_code ещё не сопоставлен активно с этой позицией
        duplicate = db.query(AgskEnstruMatch).filter(
            AgskEnstruMatch.item_id == item_id,
            AgskEnstruMatch.enstru_code == enstru_code,
            AgskEnstruMatch.is_active == True,
        ).first()
        if duplicate:
            raise ValueError(f"Код {enstru_code} уже сопоставлен с этой позицией")

        match = AgskEnstruMatch(
            agsk_code=item.code_sn or "",
            enstru_code=enstru_code,
            doc_id=item.document_id,
            item_id=item_id,
            matched_by=analyst_id,
            is_approved=False,
        )
        db.add(match)
        db.commit()
        db.refresh(match)
        return match

    def _apply_approved_match_to_all(self, db: Session, agsk_code: str, enstru_code: str) -> int:
        """Применяет утверждённое сопоставление ко всем незакрытым позициям с данным АГСК-кодом.
        Затрагивает только позиции с match_type='none' — не перезаписывает уже сопоставленные."""
        enstru_obj = db.query(Enstru).filter(Enstru.code == enstru_code).first()
        updated = db.query(PsdDocumentItem).filter(
            PsdDocumentItem.code_sn == agsk_code,
            PsdDocumentItem.match_type == "none",
        ).update({
            "enstru_code": enstru_code,
            "enstru_name": enstru_obj.name_rus if enstru_obj else None,
            "match_type": "manual",
            "match_score": 100,
            "match_reason": "Автоматически из утверждённой библиотеки сопоставлений",
        }, synchronize_session=False)
        db.commit()
        return updated

    def approve_analyst_match(self, db: Session, match_id: int, manager_id: int) -> AgskEnstruMatch:
        """Менеджер утверждает сопоставление. Синхронизирует enstru_code в позицию документа
        и автоматически применяет ко всем несопоставленным позициям с тем же АГСК-кодом."""
        match = db.query(AgskEnstruMatch).filter(AgskEnstruMatch.id == match_id).first()
        if not match:
            raise ValueError("Сопоставление не найдено")
        if not match.is_active:
            raise ValueError("Сопоставление неактивно (отклонено)")

        match.is_approved = True
        match.approved_by = manager_id
        match.approved_at = func.now()

        # Синхронизируем enstru_code в конкретную позицию документа
        if match.item_id:
            item = db.query(PsdDocumentItem).filter(PsdDocumentItem.id == match.item_id).first()
            if item:
                enstru_obj = db.query(Enstru).filter(Enstru.code == match.enstru_code).first()
                item.enstru_code = match.enstru_code
                item.enstru_name = enstru_obj.name_rus if enstru_obj else None
                item.match_type = "manual"
                item.match_score = 100
                item.match_reason = "Ручное сопоставление аналитика (утверждено менеджером)"

        db.commit()

        # Применяем утверждённую пару ко всем остальным несопоставленным позициям
        self._apply_approved_match_to_all(db, match.agsk_code, match.enstru_code)

        db.refresh(match)
        return match

    def reject_analyst_match(self, db: Session, match_id: int) -> AgskEnstruMatch:
        """Менеджер отклоняет сопоставление."""
        match = db.query(AgskEnstruMatch).filter(AgskEnstruMatch.id == match_id).first()
        if not match:
            raise ValueError("Сопоставление не найдено")
        match.is_active = False
        db.commit()
        db.refresh(match)
        return match

    def get_matches_library(
        self, db: Session,
        doc_id: Optional[int] = None,
        analyst_id: Optional[int] = None,
        date_filter: str = "all",    # "today" | "all"
        skip: int = 0,
        limit: int = 100,
    ) -> Dict:
        """Возвращает список ручных сопоставлений с их статусами (все типы вместе)."""
        query = db.query(AgskEnstruMatch)

        # Если передан конкретный документ — показываем только его
        if doc_id:
            query = query.filter(AgskEnstruMatch.doc_id == doc_id)

        if analyst_id:
            query = query.filter(AgskEnstruMatch.matched_by == analyst_id)

        if date_filter == "today":
            today_start = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0)
            query = query.filter(AgskEnstruMatch.matched_at >= today_start)

        total = query.count()
        matches = query.order_by(desc(AgskEnstruMatch.matched_at)).offset(skip).limit(limit).all()

        analyst_ids = list({m.matched_by for m in matches if m.matched_by})
        approver_ids = list({m.approved_by for m in matches if m.approved_by})
        item_ids_lib = list({m.item_id for m in matches if m.item_id})

        analysts_map = {u.id: u for u in db.query(User).filter(User.id.in_(analyst_ids)).all()} if analyst_ids else {}
        approvers_map = {u.id: u for u in db.query(User).filter(User.id.in_(approver_ids)).all()} if approver_ids else {}
        items_map = {i.id: i for i in db.query(PsdDocumentItem).filter(PsdDocumentItem.id.in_(item_ids_lib)).all()} if item_ids_lib else {}

        result = []
        for m in matches:
            if m.is_active and m.is_approved:
                status = "approved"
            elif m.is_active and not m.is_approved:
                status = "pending"
            else:
                status = "rejected"
            analyst_u = analysts_map.get(m.matched_by)
            approver_u = approvers_map.get(m.approved_by) if m.approved_by else None
            item_u = items_map.get(m.item_id) if m.item_id else None
            result.append({
                "id": m.id,
                "agsk_code": m.agsk_code,
                "enstru_code": m.enstru_code,
                "doc_id": m.doc_id,
                "item_id": m.item_id,
                "item_name": item_u.name if item_u else None,
                "matched_by": m.matched_by,
                "analyst_name": analyst_u.full_name if analyst_u else "—",
                "matched_at": m.matched_at.isoformat() if m.matched_at else None,
                "is_approved": m.is_approved,
                "is_active": m.is_active,
                "approved_by": m.approved_by,
                "approved_by_name": approver_u.full_name if approver_u else None,
                "approved_at": m.approved_at.isoformat() if m.approved_at else None,
                "status": status,
            })
        return {"items": result, "total": total}

    def parse_psd_file(self, db: Session, doc_id: int, file_path: str):
        import xml.etree.ElementTree as ET
        from collections import defaultdict

        KEYS = {
            'type': ['Тип', 'Type', 'Kind', 'Class'],
            'code_snb': ['КодСНБ', 'CodeSNB', 'Cipher', 'Code'],
            'name': ['Наименование', 'Name', 'Description'],
            'unit': ['Измеритель', 'Unit', 'MeasureUnit', 'UOM', 'Measure'],
            'vol': ['Объем', 'Quantity', 'Volume', 'Amount', 'Count'],
        }

        def _strip_ns(tree):
            for el in tree.iter():
                if '}' in el.tag:
                    el.tag = el.tag.split('}', 1)[1]
            return tree

        def _f(val):
            if not val: return 0.0
            try:
                return float(str(val).replace(',', '.').replace('\xa0', '').strip())
            except:
                return 0.0

        def _v(node, key):
            for k in KEYS.get(key, []):
                if node.get(k) is not None: return node.get(k)
            return None

        def _classify(xml_type, code_snb, unit_raw, is_resource=False, is_k2=False):
            unit = str(unit_raw).lower().strip()
            code = str(code_snb).strip()
            xt = str(xml_type).lower()

            if any(x in unit for x in ['маш.-ч', 'маш.ч', 'маш-ч', 'маш/ч', 'mach-h']): return 'SERVICES'
            if any(x in unit for x in ['чел.-ч', 'чел.ч', 'чел-ч', 'чел/ч', 'man-h']): return 'WORKS'
            if any(x in unit for x in ['т·км', 'т•км', 'ткм', 'т/км']): return 'SERVICES'
            if any(x in unit for x in ['м3 подстилающего', 'м3 основания', 'м2 поверхности', 'м2 полотна', 'переезд']): return 'WORKS'

            if code.startswith('556') or code.startswith('557'): return 'GOODS'

            digits = code.replace('-', '')
            if digits[:4].isdigit():
                prefix = int(digits[:4])
                if 1000 <= prefix <= 1999: return 'WORKS'
                if 6000 <= prefix <= 6999: return 'WORKS'
                if 3000 <= prefix <= 4999: return 'SERVICES'

            if is_k2:
                if xt in ['2', '5', '5.1']: return 'GOODS'
                if xt in ['3', '4', '6']: return 'SERVICES'
                if xt in ['1', '1.1', '1.2', '0']: return 'WORKS'
            else:
                if not is_resource:
                    if xt in ['1', '2']: return 'GOODS'
                else:
                    if xt == '2': return 'GOODS'
                    if xt == '3': return 'SERVICES'
                if xt == '6': return 'SERVICES'

            if xt in ['material', 'материал', 'equipment', 'оборудование']: return 'GOODS'
            if xt in ['machine', 'mechanism', 'механизм']: return 'SERVICES'

            if not code:
                return 'OTHER'
            return 'WORKS'

        def _cost(node, is_k2=False):
            price, total = 0.0, 0.0
            if is_k2:
                cn = node.find('Cost')
                if cn is not None:
                    price = _f(cn.get('Est_Price') or cn.get('Summary'))
                    vp = cn.find('Volume_Price')
                    total = _f(vp.get('Summary')) if vp is not None else _f(cn.get('Summary'))
            else:
                cn = node.find('СТОИМОСТЬ')
                if cn is not None:
                    tn = cn.find('ВСЕГО')
                    un = cn.find('ЕДИНИЦА')
                    if tn is not None:
                        total = _f(tn.get('Всего') or tn.get('ПЗ') or tn.get('ОТП'))
                    if un is not None:
                        price = _f(un.get('Всего') or un.get('ПЗ') or un.get('Цена'))
                if total == 0:
                    fc = node.find('Cost')
                    if fc is not None: total = _f(fc.get('total') or fc.get('Summary'))
                if price == 0: price = _f(node.get('Цена'))
            return price, total

        all_rows = []

        def _process_xml(xml_text):
            try:
                root = _strip_ns(ET.ElementTree(ET.fromstring(xml_text))).getroot()
                is_k2 = root.tag in ['LS', 'OS', 'SSR']

                official_total = 0.0
                if is_k2:
                    itn = root.find('.//Summary_LS') or root.find('.//Summary_OS') or root.find('.//SummaryDoc')
                    if itn is not None: official_total = _f(itn.get('Summary'))
                else:
                    itn = root.find('ИТОГДОК')
                    if itn is not None:
                        official_total = _f(itn.get('Всего') or itn.get('ЗатратыПодрядчика'))
                    if official_total == 0:
                        official_total = _f(root.get('Всего') or 0)

                calculated_total = 0.0
                positions = list(dict.fromkeys(
                    root.findall('.//ПОЗИЦИЯ') +
                    root.findall('.//Position') +
                    root.findall('.//Item')
                ))

                for pos in positions:
                    p_type = _v(pos, 'type') or '0'
                    p_code = _v(pos, 'code_snb') or ''
                    p_unit = _v(pos, 'unit') or ''
                    p_name = _v(pos, 'name') or ''
                    p_vol = _f(_v(pos, 'vol'))
                    item_type = _classify(p_type, p_code, p_unit, is_resource=False, is_k2=is_k2)

                    p_price, p_total_xml = _cost(pos, is_k2)
                    if p_price > 0 and p_vol > 0:
                        p_total = p_price * p_vol
                    else:
                        p_total = p_total_xml
                        if p_vol > 0 and p_total > 0: p_price = p_total / p_vol

                    extracted = 0.0
                    if item_type == 'WORKS':
                        for res in (pos.findall('РЕСУРС') + pos.findall('Resource') + pos.findall('Subitem')):
                            r_type = _classify(
                                _v(res, 'type'), _v(res, 'code_snb') or '', _v(res, 'unit') or '',
                                is_resource=True, is_k2=is_k2
                            )
                            if r_type == 'WORKS': continue
                            r_vol = _f(_v(res, 'vol'))
                            r_price, r_total_xml = _cost(res, is_k2)
                            if r_price > 0 and r_vol > 0:
                                r_total = r_price * r_vol
                            else:
                                r_total = r_total_xml
                                if r_vol > 0 and r_total > 0: r_price = r_total / r_vol
                            extracted += r_total
                            all_rows.append({
                                'name': _v(res, 'name') or '', 'code': _v(res, 'code_snb') or '',
                                'unit': _v(res, 'unit') or '', 'vol': r_vol,
                                'price': r_price, 'total': r_total, 'item_type': r_type,
                            })
                            calculated_total += r_total

                    net_total = (p_total - extracted) if item_type == 'WORKS' else p_total
                    if item_type == 'WORKS' and net_total < 0 and abs(net_total) < 5.0:
                        net_total = 0.0
                    net_price = net_total / p_vol if p_vol > 0 else p_price

                    all_rows.append({
                        'name': p_name, 'code': p_code, 'unit': p_unit,
                        'vol': p_vol, 'price': net_price, 'total': net_total, 'item_type': item_type,
                    })
                    calculated_total += net_total

                delta = official_total - calculated_total
                if abs(delta) > 1.0:
                    all_rows.append({
                        'name': 'Разница (НДС, Округления, Лимиты)',
                        'code': 'БАЛАНС', 'unit': 'компл',
                        'vol': 1.0, 'price': delta, 'total': delta, 'item_type': 'BALANCE',
                    })
            except Exception:
                pass

        if file_path.lower().endswith('.zip'):
            with zipfile.ZipFile(file_path, 'r') as z:
                for name in z.namelist():
                    if not (name.lower().endswith('.kenml') or name.lower().endswith('.xml')):
                        continue
                    try:
                        content = z.read(name)
                        for enc in ['utf-8', 'cp1251']:
                            try:
                                txt = content.decode(enc)
                                if '<' in txt[:300]:
                                    _process_xml(txt)
                                    break
                            except:
                                pass
                    except:
                        pass
        else:
            try:
                with open(file_path, 'rb') as f:
                    content = f.read()
                for enc in ['utf-8', 'cp1251']:
                    try:
                        txt = content.decode(enc)
                        if '<' in txt[:300]:
                            _process_xml(txt)
                            break
                    except:
                        pass
            except:
                pass

        db.query(PsdDocumentItem).filter(PsdDocumentItem.document_id == doc_id).delete()

        from .psd_analyzer.analyzer import clean_product_name

        grouped = {}
        for row in all_rows:
            itype = row['item_type']
            code = str(row['code']).strip()
            name = str(row['name']).strip()
            if itype == 'BALANCE':
                key = f"BALANCE_{name}"
                grouped[key] = row
                continue
            key = (itype, code, name)
            if key not in grouped:
                grouped[key] = {**row, 'total': 0.0, 'vol': 0.0}
            grouped[key]['total'] += float(row['total'])
            grouped[key]['vol'] += float(row['vol'])

        final_rows = []
        for row in grouped.values():
            vol = row['vol']
            total = row['total']
            row['price'] = total / vol if vol > 0 else 0.0
            final_rows.append(row)

        for idx, row in enumerate(final_rows, 1):
            name = str(row['name']).strip() or f'Позиция {idx}'
            code = str(row['code']).strip()
            itype = row['item_type']
            db.add(PsdDocumentItem(
                document_id=doc_id,
                position_number=str(idx),
                name=name,
                code_sn=code if code and code != 'БАЛАНС' else None,
                unit=str(row['unit']),
                volume=row['vol'],
                price=float(row['price']),
                total_amount=row['total'],
                clean_name=clean_product_name(name),
                is_product=(itype == 'GOODS'),
                item_type=itype,
                match_type='none',
            ))

        db.commit()

        doc = db.query(ExternalDocument).filter(ExternalDocument.id == doc_id).first()
        if doc and doc.status != 'ASSIGNED_TO_ANALYST':
            doc.status = 'PARSED'
            db.commit()

        self._run_auto_matching_for_document(db, doc_id)

    def parse_smeta_file(self, db: Session, doc_id: int, file_path: str):
        import openpyxl
        from .importers.excel_importer import extract_code

        wb = openpyxl.load_workbook(file_path, data_only=True)
        ws = wb["Позиции для загрузки"] if "Позиции для загрузки" in wb.sheetnames else wb.active
        db.query(PsdDocumentItem).filter(PsdDocumentItem.document_id == doc_id).delete()

        position_idx = 1
        for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            if not any(cell is not None and str(cell).strip() for cell in row):
                continue
            row_data = list(row) + [None] * max(0, 17 - len(row))
            enstru_code = str(row_data[1]).strip() if row_data[1] else None
            name = str(row_data[3]).strip() if row_data[3] else None
            if not name:
                name = str(row_data[2]).strip() if row_data[2] else f"Позиция {position_idx}"
            unit = str(row_data[5]).strip() if row_data[5] else ""
            agsk_code = str(row_data[13]).strip() if row_data[13] else None
            if agsk_code and agsk_code.lower() == "прайс-лист":
                agsk_code = None
            try:
                volume = float(row_data[6]) if row_data[6] is not None else 0.0
            except (ValueError, TypeError):
                volume = 0.0
            try:
                price = float(row_data[7]) if row_data[7] is not None else 0.0
            except (ValueError, TypeError):
                price = 0.0
            try:
                total_amount = float(row_data[8]) if row_data[8] is not None else (volume * price)
            except (ValueError, TypeError):
                total_amount = volume * price
            if not enstru_code and not name:
                continue
            item = PsdDocumentItem(
                document_id=doc_id, position_number=str(position_idx), name=name,
                code_sn=agsk_code, unit=unit, volume=volume, price=price,
                total_amount=total_amount,
                clean_name=name.split('/')[0].strip() if '/' in name else name,
                is_product=True,
                enstru_code=enstru_code,
                enstru_name=str(row_data[2]).strip() if row_data[2] else None,
                match_type="auto" if enstru_code else "none",
                match_score=100.0 if enstru_code else None,
                match_reason="Из файла сметы" if enstru_code else None,
            )
            db.add(item)
            position_idx += 1

        db.commit()
        doc = db.query(ExternalDocument).filter(ExternalDocument.id == doc_id).first()
        if doc and doc.status != "ASSIGNED_TO_ANALYST":
            doc.status = "PARSED"
            db.commit()

    def export_matches_to_excel(self, db: Session, format_type: str = "full"):
        matches = db.query(AgskEnstruMatch).filter(
            AgskEnstruMatch.is_active == True,
            AgskEnstruMatch.is_approved == True,
        ).order_by(AgskEnstruMatch.agsk_code).all()
        data = [{
            "Код АГСК": m.agsk_code,
            "Код ЕНС ТРУ": m.enstru_code,
            "Дата сопоставления": m.matched_at.strftime('%d.%m.%Y') if m.matched_at else None,
            "Дата утверждения": m.approved_at.strftime('%d.%m.%Y') if m.approved_at else None,
        } for m in matches]
        path = f"/tmp/export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        pd.DataFrame(data).to_excel(path, index=False)
        return {"file_path": path}

    def export_full_analysis_report(self, db: Session, doc_id: int) -> Optional[str]:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        q = db.query(PsdDocumentItem, Agsk).outerjoin(Agsk, PsdDocumentItem.code_sn == Agsk.code)
        q = q.filter(PsdDocumentItem.document_id == doc_id)
        items_data = q.order_by(PsdDocumentItem.id).all()

        # Deduplicate by item.id — outerjoin with Agsk can produce duplicate rows
        # if code_sn matches multiple Agsk records, which inflates averages
        _seen_item_ids: set = set()
        _deduped: list = []
        for _it, _ag in items_data:
            if _it.id not in _seen_item_ids:
                _seen_item_ids.add(_it.id)
                _deduped.append((_it, _ag))
        items_data = _deduped

        if not items_data:
            return None

        doc = db.query(ExternalDocument).filter(ExternalDocument.id == doc_id).first()
        agsk_codes = list({item.code_sn for item, _ in items_data if item.code_sn})
        enstru_codes_set = list({item.enstru_code for item, _ in items_data if item.enstru_code})

        direct_ktp_map, group_ktp_map, suppliers_by_enstru = PsdAnalystService._build_dvc_maps(
            db, agsk_codes, enstru_codes_set
        )

        wb = openpyxl.Workbook()
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="1B5E20", end_color="1B5E20", fill_type="solid")
        sub_header_fill = PatternFill(start_color="E8F5E9", end_color="E8F5E9", fill_type="solid")
        border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )

        type_letters = {'GOODS': 'Т', 'WORKS': 'Р', 'SERVICES': 'У', 'OTHER': 'П', 'BALANCE': 'Б'}
        type_row_fills = {
            'WORKS': PatternFill(start_color="EBF5FB", end_color="EBF5FB", fill_type="solid"),
            'SERVICES': PatternFill(start_color="EAF7EA", end_color="EAF7EA", fill_type="solid"),
            'OTHER': PatternFill(start_color="F5F5F5", end_color="F5F5F5", fill_type="solid"),
            'BALANCE': PatternFill(start_color="FEF9E7", end_color="FEF9E7", fill_type="solid"),
        }

        # Группируем позиции по типу
        grouped = defaultdict(list)
        for item, agsk in items_data:
            itype = item.item_type or 'GOODS'
            grouped[itype].append((item, agsk))

        # ── ЛИСТ 1: СМЕТА ────────────────────────────────────────────────────
        ws1 = wb.active
        ws1.title = "Смета"

        ws1.merge_cells('A1:K1')
        ws1['A1'] = "СМЕТА ЗАКУПОК"
        ws1['A1'].font = Font(size=16, bold=True)
        ws1['A1'].alignment = Alignment(horizontal='center')
        if doc:
            ws1.merge_cells('A2:K2')
            ws1['A2'] = f"Наименование проекта: {doc.bank_name}"
            ws1['A2'].font = Font(bold=True, size=12)

        columns1 = [
            "№", "Код ЕНС ТРУ", "Наименование закупаемых товаров услуг работ",
            "Единица измерения(МКЕИ)", "Количество, объём", "Цена за единицу тенге(без НДС)",
            "Сумма планируемая для закупок ТРУ", "КОД АГСК", "КТП", "ВЦ %", "Сумма ВЦ тенге без НДС"
        ]

        def write_header(ws, row):
            for col_idx, col_name in enumerate(columns1, 1):
                cell = ws.cell(row=row, column=col_idx, value=col_name)
                cell.font = header_font
                cell.fill = header_fill
                cell.border = border
                cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
            ws.row_dimensions[row].height = 45
            return row + 1

        def fill_section(ws, curr_row, section_title, section_items, is_first):
            if not section_items:
                return curr_row, Decimal('0'), Decimal('0'), Decimal('0'), 0

            if is_first:
                curr_row = write_header(ws, curr_row)

            # Заголовок секции
            ws.merge_cells(f'A{curr_row}:K{curr_row}')
            c = ws.cell(row=curr_row, column=1, value=section_title)
            c.font = Font(bold=True, size=12)
            c.fill = sub_header_fill
            curr_row += 1

            section_total = Decimal('0')
            section_vc = Decimal('0')
            dvc_pct_sum = Decimal('0')
            dvc_pct_count = 0

            for idx, (item, agsk) in enumerate(section_items, 1):
                itype = item.item_type or 'GOODS'
                num_label = f"{idx} {type_letters.get(itype, 'П')}"
                dvc_p = PsdAnalystService._calc_min_dvc(item, direct_ktp_map, group_ktp_map, suppliers_by_enstru)
                i_total = Decimal(str(item.total_amount or 0))
                vc_a = i_total * (dvc_p / 100)
                section_total += i_total
                section_vc += vc_a
                dvc_pct_sum += dvc_p
                dvc_pct_count += 1

                row_data = [
                    num_label,
                    item.enstru_code or "",
                    item.name or "",
                    item.unit or "",
                    float(item.volume or 0),
                    float(item.price or 0),
                    float(i_total),
                    item.code_sn or "",
                    "Да" if item.match_type != 'none' else ("100%" if itype in ('WORKS', 'SERVICES') else "Нет"),
                    float(dvc_p),
                    float(vc_a)
                ]

                row_fill = type_row_fills.get(itype)
                for c_idx, val in enumerate(row_data, 1):
                    cell = ws.cell(row=curr_row, column=c_idx, value=val)
                    cell.border = border
                    if c_idx in [5, 6, 7, 10, 11]:
                        cell.number_format = '#,##0.00'
                    if row_fill:
                        cell.fill = row_fill
                curr_row += 1

            # Итог секции — простое арифметическое среднее ВЦ% (все позиции, включая 0)
            avg_dvc_sec = dvc_pct_sum / dvc_pct_count if dvc_pct_count > 0 else Decimal('0')
            ws.merge_cells(f'A{curr_row}:F{curr_row}')
            c = ws.cell(row=curr_row, column=1, value=f"Итого по {section_title.lower()}:")
            c.font = Font(bold=True)
            c.alignment = Alignment(horizontal='right')
            c7 = ws.cell(row=curr_row, column=7, value=float(section_total))
            c7.font = Font(bold=True)
            c7.number_format = '#,##0.00'
            c10 = ws.cell(row=curr_row, column=10, value=float(avg_dvc_sec))
            c10.font = Font(bold=True)
            c10.number_format = '0.00'
            c11 = ws.cell(row=curr_row, column=11, value=float(section_vc))
            c11.font = Font(bold=True)
            c11.number_format = '#,##0.00'
            curr_row += 2  # отступ после секции

            return curr_row, section_total, section_vc, dvc_pct_sum, dvc_pct_count

        curr_row = 6
        total_sum = Decimal('0')
        total_vc_sum = Decimal('0')
        total_dvc_pct_sum = Decimal('0')
        total_dvc_pct_count = 0

        section_configs = [
            ('1. Товары (Т)', 'GOODS'),
            ('2. Работы (Р)', 'WORKS'),
            ('3. Услуги (У)', 'SERVICES'),
            ('4. Прочее (П)', 'OTHER'),
        ]

        # Build item_id → num_label map so sheet 2 uses the same numbering as sheet 1
        item_num_labels: dict = {}
        for _, itype_key in section_configs:
            tl = type_letters.get(itype_key, 'П')
            for sec_idx, (sec_item, _) in enumerate(grouped.get(itype_key, []), 1):
                item_num_labels[sec_item.id] = f"{sec_idx} {tl}"

        is_first = True
        for section_title, itype in section_configs:
            items_in_section = grouped.get(itype, [])
            if not items_in_section:
                continue
            curr_row, sec_total, sec_vc, sec_dvc_sum, sec_dvc_count = fill_section(
                ws1, curr_row, section_title, items_in_section, is_first
            )
            total_sum += sec_total
            total_vc_sum += sec_vc
            total_dvc_pct_sum += sec_dvc_sum
            total_dvc_pct_count += sec_dvc_count
            is_first = False

        # Общий итог — простое арифметическое среднее ВЦ% по всем позициям
        avg_total_dvc = total_dvc_pct_sum / total_dvc_pct_count if total_dvc_pct_count > 0 else Decimal('0')
        ws1.merge_cells(f'A{curr_row}:F{curr_row}')
        ws1.cell(row=curr_row, column=1, value="ИТОГО ПО СМЕТЕ:").font = Font(bold=True, size=11)
        ws1.cell(row=curr_row, column=1).alignment = Alignment(horizontal='right')
        ws1.cell(row=curr_row, column=7, value=float(total_sum)).font = Font(bold=True, size=11)
        ws1.cell(row=curr_row, column=7).number_format = '#,##0.00'
        c10_total = ws1.cell(row=curr_row, column=10, value=float(avg_total_dvc))
        c10_total.font = Font(bold=True, size=11)
        c10_total.number_format = '0.00'
        ws1.cell(row=curr_row, column=11, value=float(total_vc_sum)).font = Font(bold=True, size=11)
        ws1.cell(row=curr_row, column=11).number_format = '#,##0.00'

        # ── ЛИСТ 2: АНАЛИЗ ПСД — только GOODS с сопоставлением ──────────────
        ws2 = wb.create_sheet("Анализ ПСД")
        columns2 = [
            "№ позиции", "Наименование позиции", "Код АГСК", "Полное название АГСК",
            "Ед. изм.", "Объем", "Цена", "Сумма",
            "Источник", "Производитель из РКТП", "БИН",
            "Наименование товара из РКТП", "ВЦ%", "Адрес", "Регион",
            "Коды АГСК из Реестра КТП"
        ]
        for col_idx, col_name in enumerate(columns2, 1):
            cell = ws2.cell(row=1, column=col_idx, value=col_name)
            cell.font = header_font
            cell.fill = PatternFill("solid", fgColor="2C3E50")
            cell.border = border
            cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

        FILL_MAP = {
            "Выбор аналитика": PatternFill("solid", fgColor="D5F5E3"),
            "Авто-подбор (КТП)": PatternFill("solid", fgColor="D6EAF8"),
            "Реестр (по коду ЕНС)": PatternFill("solid", fgColor="FEF9E7"),
        }

        curr_row2 = 2

        # Только GOODS с сопоставлением (есть enstru_code)
        for item, agsk in grouped.get('GOODS', []):
            if not item.enstru_code:
                continue  # без сопоставления — пропускаем

            num_label = item_num_labels.get(item.id, "? Т")
            base = [
                num_label,
                item.name,
                item.code_sn,
                agsk.full_name if agsk else "",
                item.unit,
                float(item.volume or 0),
                float(item.price or 0),
                float(item.total_amount or 0),
            ]

            rows_to_add = []
            selected_ktp_ids = set()

            auto_ktp = direct_ktp_map.get(item.code_sn) or group_ktp_map.get(item.code_sn)
            if auto_ktp and auto_ktp.id not in selected_ktp_ids:
                dvc = float(
                    re.sub(r'[^0-9.]', '', auto_ktp.dvc_percent).replace(',', '.')) if auto_ktp.dvc_percent else 0
                if dvc <= 0:  # добавь
                    selected_ktp_ids.add(auto_ktp.id)
                else:
                    rows_to_add.append(base + [
                        "Авто-подбор (КТП)",
                        auto_ktp.company_name, auto_ktp.bin_iin, auto_ktp.product_name,
                        dvc, auto_ktp.production_address, auto_ktp.region_kato,
                        ", ".join(auto_ktp.agsk3_codes) if auto_ktp.agsk3_codes else "",
                    ])
                    selected_ktp_ids.add(auto_ktp.id)

            all_suppliers = suppliers_by_enstru.get(item.enstru_code, [])
            for s in all_suppliers:
                if s.id not in selected_ktp_ids:
                    dvc = float(re.sub(r'[^0-9.]', '', s.dvc_percent).replace(',', '.')) if s.dvc_percent else 0
                    if dvc <= 0:
                        continue
                    rows_to_add.append(base + [
                        "Реестр (по коду ЕНС)",
                        s.company_name, s.bin_iin, s.product_name,
                        dvc, s.production_address, s.region_kato,
                        ", ".join(s.agsk3_codes) if s.agsk3_codes else "",
                    ])

            if not rows_to_add:
                continue  # сопоставление есть (enstru_code), но КТП не найден — пропускаем

            for r_data in rows_to_add:
                src = r_data[8]
                fill = FILL_MAP.get(src, PatternFill())
                for c_idx, val in enumerate(r_data, 1):
                    cell = ws2.cell(row=curr_row2, column=c_idx, value=val)
                    cell.fill = fill
                    cell.border = border
                    if c_idx in [6, 7, 8, 13]:
                        cell.number_format = '#,##0.00'
                ws2.cell(row=curr_row2, column=12).font = Font(bold=True)
                curr_row2 += 1

        # Автоширина колонок
        for ws in [ws1, ws2]:
            for i, col in enumerate(ws.columns, 1):
                m_l = 0
                for cell in col:
                    try:
                        if cell.value and len(str(cell.value)) > m_l:
                            m_l = len(str(cell.value))
                    except:
                        pass
                ws.column_dimensions[get_column_letter(i)].width = min(m_l + 2, 40)

        os.makedirs("/tmp", exist_ok=True)
        file_path = f"/tmp/psd_analysis_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        wb.save(file_path)
        return file_path

    def generate_psd_conclusion_docx(self, db: Session, doc_id: int, current_user: User) -> Optional[str]:
        doc_entity = db.query(ExternalDocument).filter(ExternalDocument.id == doc_id).first()
        if not doc_entity:
            return None

        total_items = db.query(PsdDocumentItem).filter(PsdDocumentItem.document_id == doc_id).count()
        matched_items = db.query(PsdDocumentItem).filter(
            PsdDocumentItem.document_id == doc_id,
            PsdDocumentItem.match_type != 'none'
        ).count()
        total_amount = db.query(func.sum(PsdDocumentItem.total_amount)).filter(
            PsdDocumentItem.document_id == doc_id
        ).scalar() or 0

        items = db.query(PsdDocumentItem).filter(PsdDocumentItem.document_id == doc_id).all()

        agsk_codes_doc = list({it.code_sn for it in items if it.code_sn})
        enstru_codes_doc = list({it.enstru_code for it in items if it.enstru_code})
        direct_ktp_map, group_ktp_map, suppliers_map = PsdAnalystService._build_dvc_maps(
            db, agsk_codes_doc, enstru_codes_doc
        )

        goods_amount = Decimal('0')
        works_amount = Decimal('0')
        services_amount = Decimal('0')
        other_amount = Decimal('0')
        goods_vc = Decimal('0')
        works_vc = Decimal('0')
        services_vc = Decimal('0')
        total_vc_amount = Decimal('0')
        goods_dvc_sum = Decimal('0'); goods_dvc_count = 0
        works_dvc_sum = Decimal('0'); works_dvc_count = 0
        services_dvc_sum = Decimal('0'); services_dvc_count = 0
        total_dvc_sum = Decimal('0'); total_dvc_count = 0

        for it in items:
            itype = it.item_type or 'GOODS'
            i_total = Decimal(str(it.total_amount or 0))
            dvc = PsdAnalystService._calc_min_dvc(it, direct_ktp_map, group_ktp_map, suppliers_map)
            vc_part = i_total * (dvc / 100)
            total_dvc_sum += dvc
            total_dvc_count += 1
            if itype == 'GOODS':
                goods_amount += i_total
                goods_vc += vc_part
                goods_dvc_sum += dvc
                goods_dvc_count += 1
            elif itype == 'WORKS':
                works_amount += i_total
                works_vc += vc_part
                works_dvc_sum += dvc
                works_dvc_count += 1
            elif itype == 'SERVICES':
                services_amount += i_total
                services_vc += vc_part
                services_dvc_sum += dvc
                services_dvc_count += 1
            else:
                other_amount += i_total
            total_vc_amount += vc_part

        # Простое арифметическое среднее (все позиции включая 0)
        avg_vc_percent = total_dvc_sum / total_dvc_count if total_dvc_count > 0 else Decimal('0')
        goods_avg_dvc = goods_dvc_sum / goods_dvc_count if goods_dvc_count > 0 else Decimal('0')
        works_avg_dvc = works_dvc_sum / works_dvc_count if works_dvc_count > 0 else Decimal('0')
        services_avg_dvc = services_dvc_sum / services_dvc_count if services_dvc_count > 0 else Decimal('0')

        doc = Document()
        title = doc.add_heading('ЗАКЛЮЧЕНИЕ АНАЛИТИКА ДРВЦ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        p.add_run('Дата формирования: ').bold = True
        p.add_run(datetime.now().strftime('%d.%m.%Y %H:%M'))

        p = doc.add_paragraph()
        p.add_run('Аналитик: ').bold = True
        p.add_run(f"{current_user.full_name}")

        p = doc.add_paragraph()
        p.add_run('Наименование Банка: ').bold = True
        p.add_run(f"{doc_entity.bank_name}")

        doc.add_heading('Результаты анализа', level=1)

        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Показатель'
        hdr_cells[1].text = 'Значение'

        data = [
            ('Общее количество позиций', str(total_items)),
            ('Сопоставлено позиций (Товары)', str(matched_items)),
            ('Процент сопоставления', f"{(matched_items / total_items * 100):.1f}%" if total_items > 0 else "0%"),
            ('', ''),
            ('Сумма Товары (Т)', f"{float(goods_amount):,.2f} тенге"),
            ('Средний ВЦ% по Товарам (Т)', f"{float(goods_avg_dvc):.2f}%"),
            ('Сумма Работы (Р)', f"{float(works_amount):,.2f} тенге"),
            ('Средний ВЦ% по Работам (Р)', f"{float(works_avg_dvc):.2f}%"),
            ('Сумма Услуги (У)', f"{float(services_amount):,.2f} тенге"),
            ('Средний ВЦ% по Услугам (У)', f"{float(services_avg_dvc):.2f}%"),
        ]
        if other_amount > 0:
            data.append(('Сумма Прочее (П)', f"{float(other_amount):,.2f} тенге"))
        data += [
            ('Общая сумма проекта', f"{float(total_amount):,.2f} тенге"),
            ('', ''),
            ('Прогнозируемая сумма ВЦ', f"{float(total_vc_amount):,.2f} тенге"),
            ('Средний ВЦ% (общий)', f"{float(avg_vc_percent):.2f}%"),
        ]

        for label, value in data:
            row_cells = table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = value

        doc.add_paragraph()
        doc.add_heading('Выводы', level=1)
        conclusion_text = (
            f"В ходе анализа проектно-сметной документации '{doc_entity.bank_name}' "
            f"было обработано {total_items} позиций. "
            f"Уровень внутристрановой ценности (ВЦ) по проекту оценивается в {float(avg_vc_percent):.2f}%. "
        )
        doc.add_paragraph(conclusion_text)

        if doc_entity.analyst_comment:
            doc.add_paragraph()
            doc.add_heading('Дополнительный комментарий аналитика', level=1)
            doc.add_paragraph(doc_entity.analyst_comment)

        doc.add_paragraph().add_run('\n\n__________________________ / Подпись /').italic = True

        os.makedirs("/tmp", exist_ok=True)
        file_path = f"/tmp/conclusion_{doc_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(file_path)
        return file_path
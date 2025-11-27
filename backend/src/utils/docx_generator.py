from docx import Document
from docx.shared import Pt
from io import BytesIO


def generate_docx(application) -> bytes:
    doc = Document()
    doc.add_heading("Заявка на финансирование", 0)

    p = doc.add_paragraph()
    p.add_run("№ заявки: ").bold = True
    p.add_run(str(application.number))

    fields = [
        ("Вид потребности", application.need_type),
        ("Код ЕНС ТРУ", application.enstru_code),
        ("Наименование", application.enstru_name),
        ("Доп. характеристика", application.additional_specs),
        ("Источник финансирования", application.funding_source),
        ("Количество", str(application.quantity)),
        ("Цена без НДС", f"{application.marketing_price or 0:,.2f} ₸"),
        ("Сумма без НДС", f"{application.planned_total_no_vat or 0:,.2f} ₸"),
    ]

    for title, value in fields:
        if value:
            p = doc.add_paragraph()
            p.add_run(f"{title}: ").bold = True
            p.add_run(str(value))

    doc.add_paragraph(f"\nСтатус: {application.state.value.upper()}")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()